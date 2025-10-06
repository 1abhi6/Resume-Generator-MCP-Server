from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, uuid


def add_separator(document):
    """Add a simple horizontal separator."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run("─" * 60)
    run.font.color.rgb = RGBColor(180, 180, 180)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section_heading(document, title):
    """Stylish section heading."""
    heading = document.add_paragraph()
    run = heading.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(60, 60, 60)
    document.add_paragraph("")  # spacing
    add_separator(document)
    document.add_paragraph("")  # spacing


def generate_resume_docx(profile_data):
    document = Document()

    # Set margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # ===== HEADER =====
    header = document.add_paragraph()
    header_run = header.add_run(profile_data["name"].upper())
    header_run.bold = True
    header_run.font.size = Pt(20)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_header = document.add_paragraph()
    contact_line = f"{profile_data.get('title', '')}"
    if profile_data.get("location"):
        contact_line += f" | {profile_data['location']}"
    if profile_data.get("email"):
        contact_line += f" | {profile_data['email']}"
    if profile_data.get("linkedin"):
        contact_line += f" | {profile_data['linkedin']}"

    sub_run = sub_header.add_run(contact_line)
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)
    sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_separator(document)
    document.add_paragraph("")

    # ===== SUMMARY =====
    if profile_data.get("summary"):
        add_section_heading(document, "Summary")
        document.add_paragraph(profile_data["summary"])

    # ===== EXPERIENCE =====
    if profile_data.get("experience"):
        add_section_heading(document, "Experience")
        for exp in profile_data["experience"]:
            p = document.add_paragraph(exp, style="List Bullet")
            p_format = p.paragraph_format
            p_format.space_after = Pt(2)

    # ===== EDUCATION =====
    if profile_data.get("education"):
        add_section_heading(document, "Education")
        for edu in profile_data["education"]:
            p = document.add_paragraph(edu, style="List Bullet")
            p.paragraph_format.space_after = Pt(2)

    # ===== SKILLS =====
    if profile_data.get("skills"):
        add_section_heading(document, "Skills")
        skills_line = " · ".join(profile_data["skills"])
        p = document.add_paragraph(skills_line)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ===== SAVE =====
    file_name = f"resume_{uuid.uuid4().hex}.docx"
    document.save(file_name)
    # return {"docx_path": file_name}


profile_data = {
    "name": "Abhishek Gupta",
    "title": "Agentic AI Developer",
    "summary": "Skilled in LangGraph, CrewAI, and FastMCP, focusing on building intelligent automation tools.",
    "experience": [
        "Developed MCP servers for AI workflows.",
        "Integrated FastAPI with LangChain and S3 for scalable automation.",
        "Built and deployed multi-agent architectures for startups.",
    ],
    "education": ["B.Tech in Computer Science - Pune Institute of Technology (2022)"],
    "skills": ["LangChain", "FastMCP", "CrewAI", "AWS", "Python", "FastAPI"],
}


generate_resume_docx(profile_data)
